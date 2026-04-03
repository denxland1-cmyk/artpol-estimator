"""
ARTPOL — Генератор договора подряда
Использует python-docx напрямую. Без внешних скриптов.
Открывает шаблон .docx, подставляет данные, сохраняет.
"""

import os
import logging
from pathlib import Path
from datetime import datetime, timezone, timedelta
from copy import deepcopy

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

logger = logging.getLogger(__name__)

TEMPLATE_PATH = Path(__file__).parent / "ШАБЛОН_ДОГОВОРА_НОВЫИ_.docx"
MSK = timezone(timedelta(hours=3))

MONTHS_RU = {
    1: "января", 2: "февраля", 3: "марта", 4: "апреля",
    5: "мая", 6: "июня", 7: "июля", 8: "августа",
    9: "сентября", 10: "октября", 11: "ноября", 12: "декабря",
}


def _num_to_words(n: int) -> str:
    """Число → прописью (до миллиона)."""
    if n == 0:
        return "ноль"

    ones = ["", "одна", "две", "три", "четыре", "пять", "шесть",
            "семь", "восемь", "девять", "десять", "одиннадцать",
            "двенадцать", "тринадцать", "четырнадцать", "пятнадцать",
            "шестнадцать", "семнадцать", "восемнадцать", "девятнадцать"]
    tens = ["", "", "двадцать", "тридцать", "сорок", "пятьдесят",
            "шестьдесят", "семьдесят", "восемьдесят", "девяносто"]
    hundreds = ["", "сто", "двести", "триста", "четыреста", "пятьсот",
                "шестьсот", "семьсот", "восемьсот", "девятьсот"]

    parts = []
    orig = n

    if n >= 1000:
        t = n // 1000
        n = n % 1000
        if t >= 100:
            parts.append(hundreds[t // 100])
            t = t % 100
        if t >= 20:
            parts.append(tens[t // 10])
            t = t % 10
        if t > 0:
            if t == 1:
                parts.append("одна")
            elif t == 2:
                parts.append("две")
            else:
                parts.append(ones[t])

        # тысяча/тысячи/тысяч
        tt = (orig // 1000) % 100
        last_digit = (orig // 1000) % 10
        if 11 <= tt <= 19:
            parts.append("тысяч")
        elif last_digit == 1:
            parts.append("тысяча")
        elif 2 <= last_digit <= 4:
            parts.append("тысячи")
        else:
            parts.append("тысяч")

    if n >= 100:
        parts.append(hundreds[n // 100])
        n = n % 100
    if n >= 20:
        parts.append(tens[n // 10])
        n = n % 10
    if n > 0:
        parts.append(ones[n])

    return " ".join(p for p in parts if p)


def _replace_in_paragraph(paragraph, old_text, new_text):
    """
    Заменяет текст в параграфе, даже если он разбит на несколько runs.
    """
    for run in paragraph.runs:
        if old_text in run.text:
            run.text = run.text.replace(old_text, new_text)
            return True

    full_text = paragraph.text
    if old_text not in full_text:
        return False

    runs_with_text = [(i, run) for i, run in enumerate(paragraph.runs) if run.text]
    if not runs_with_text:
        return False

    concat = ""
    for i, run in runs_with_text:
        concat += run.text

    find_pos = concat.find(old_text)
    if find_pos == -1:
        return False

    find_end = find_pos + len(old_text)
    new_concat = concat[:find_pos] + new_text + concat[find_end:]

    first_run_idx = runs_with_text[0][0]
    paragraph.runs[first_run_idx].text = new_concat
    for i, run in runs_with_text[1:]:
        run.text = ""

    return True


def _set_table_cell(cell, text, bold=False, size=9, align=WD_ALIGN_PARAGRAPH.LEFT):
    """Настраивает ячейку таблицы сметы."""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    run.bold = bold


def _shade_cells(row, color):
    """Заливает ячейки строки цветом."""
    for cell in row.cells:
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
        cell._tc.get_or_add_tcPr().append(shading)


def _insert_estimate_table(doc, parsed, estimate, area, thickness, grade, include_sand_removal):
    """Вставляет таблицу сметы после п.2.1."""
    # Находим параграф с суммой
    target_idx = None
    for i, para in enumerate(doc.paragraphs):
        if "рублей 00 копеек" in para.text:
            target_idx = i
            break

    if target_idx is None:
        logger.warning("Не нашёл п.2.1 для вставки таблицы")
        return

    # Если данные из АМО (без детализации) — пропускаем таблицу
    s = estimate["sand"]
    if s.get("sand_tons", 0) == 0 and s.get("sand_cost", 0) == 0:
        return
    c = estimate["cement"]
    f = estimate["fiber"]
    fl = estimate["film"]
    iz = estimate["izoflex"]
    eq = estimate["equipment_delivery"]
    w = estimate["work"]
    k = estimate.get("keramzit")
    volume = round(area * thickness / 1000, 3)

    ker_data = parsed.get("keramzit") or {}
    ker_area = ker_data.get("area_m2", 0)
    ker_thick = ker_data.get("thickness_mm", 0)

    L = WD_ALIGN_PARAGRAPH.LEFT
    R = WD_ALIGN_PARAGRAPH.RIGHT
    C = WD_ALIGN_PARAGRAPH.CENTER
    BLUE_LIGHT = "DCE6F1"
    BLUE_HEADER = "B8CCE4"

    # Создаём таблицу
    table = doc.add_table(rows=0, cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Ширина на всю страницу
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>')
    for old in tblPr.findall(qn("w:tblW")):
        tblPr.remove(old)
    tblPr.append(tblW)

    def add_row(data, shade=None):
        row = table.add_row()
        for i, (text, bold, align) in enumerate(data):
            _set_table_cell(row.cells[i], text, bold=bold, size=9, align=align)
        if shade:
            _shade_cells(row, shade)
        return row

    # Шапка: керамзит (если есть)
    if k and ker_area:
        add_row([("", False, L), (f"{ker_area}", False, R), ("", False, L),
                 (f"{ker_thick}", False, R), ("", False, R)], shade=BLUE_LIGHT)

    # Шапка: стяжка
    add_row([("", False, L), (f"{area}", False, R), ("", False, L),
             (f"{thickness}", False, R), (f"{volume}", False, R)], shade=BLUE_LIGHT)

    # Заголовки
    add_row([("материалы и транспортные расходы", True, L), ("ед. изм.", True, C),
             ("кол-во ед.", True, C), ("Стоимость ед.", True, C),
             ("ИТОГО", True, R)], shade=BLUE_HEADER)

    # Песок
    add_row([(f"Песок {s['sand_tons']}т+доставка", False, L), ("рейс", False, C),
             ("1", False, R), (f"{s['total']}", False, R), (f"{s['total']}", False, R)])

    # Цемент
    ppb = c["cement_cost"] // c["bags"] if c["bags"] > 0 else 0
    add_row([("Цемент М500 по 50 кг", False, L), ("мешок", False, C),
             (f"{c['bags']}", False, R), (f"{ppb}", False, R), (f"{c['cement_cost']}", False, R)])

    # Фибра
    add_row([("Фибра ВСМ 12 мм 20 мк", False, L), ("кг", False, C),
             (f"{f['kg']}", False, R), ("300", False, R), (f"{f['cost']}", False, R)])

    # Плёнка
    add_row([("Пленка 60 мкр техническая", False, L), ("м2", False, C),
             (f"{fl['m2']}", False, R), ("10", False, R), (f"{fl['cost']}", False, R)])

    # Izoflex
    add_row([("IZOFLEX 10 мм", False, L), ("пог.м.", False, C),
             (f"{iz['meters']}", False, R), ("20", False, R), (f"{iz['cost']}", False, R)])

    # Керамзит материалы
    if k:
        add_row([("Керамзит (0,075м3)", False, L), ("мешок", False, C),
                 (f"{k['keramzit_bags']}", False, R), ("340", False, R),
                 (f"{k['keramzit_cost']}", False, R)])
        add_row([("Армированная пленка 100г/м", False, L), ("м2", False, C),
                 (f"{k['reinforced_film_m2']}", False, R), ("40", False, R),
                 (f"{k['reinforced_film_cost']}", False, R)])
        add_row([("Металлическая сетка", False, L), ("м2", False, C),
                 (f"{k['mesh_m2']}", False, R), ("120", False, R),
                 (f"{k['mesh_cost']}", False, R)])

    # Доставка материалов
    add_row([("Доставка материалов", False, L), ("рейс", False, C),
             ("1", False, R), (f"{c['delivery']}", False, R), (f"{c['delivery']}", False, R)])

    # Доставка оборудования
    add_row([("Доставка/вывоз оборудования", False, L), ("рейс", False, C),
             ("1", False, R), (f"{eq['cost']}", False, R), (f"{eq['cost']}", False, R)])

    # *Работы
    add_row([("*Работы", True, L), ("", False, C), ("", False, R),
             ("руб./м2", False, R), ("", False, R)], shade=BLUE_HEADER)

    # Вывоз песка
    if include_sand_removal:
        add_row([("Вывоз\\довоз песка", False, L), ("", False, C),
                 ("", False, R), ("", False, R), ("5000", False, R)])

    # Работа керамзит
    if k:
        add_row([("Устройство керамзитного основания", False, L), ("м2", False, C),
                 (f"{ker_area}", False, R), (f"{k['keramzit_work_rate']}", False, R),
                 (f"{k['keramzit_work_cost']}", False, R)])

    # Работа стяжка
    work_rate = w.get("rate", "").replace("₽", "")
    add_row([("Устройство полусухой стяжки пола", False, L), ("м2", False, C),
             (f"{area}", False, R), (work_rate, False, R), (f"{w['cost']}", False, R)])

    # ИТОГО
    grand = estimate["grand_total"]
    if include_sand_removal:
        grand += 5000
    add_row([("", False, L), ("", False, C), ("", False, R),
             ("ИТОГО стяжка", True, R), (f"{grand}", True, R)])

    # Перемещаем таблицу после п.2.1
    # python-docx добавляет таблицу в конец, нужно переместить
    target_para = doc.paragraphs[target_idx]
    target_para._p.addnext(table._tbl)


def _replace_in_doc(doc, old_text, new_text):
    """Заменяет текст во всём документе (параграфы + таблицы)."""
    for para in doc.paragraphs:
        _replace_in_paragraph(para, old_text, new_text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_paragraph(para, old_text, new_text)


def generate_contract(
    parsed: dict,
    estimate: dict,
    client_data: dict,
    grade: str = "М150",
    include_sand_removal: bool = False,
    output_path: str = None,
) -> str:
    """
    Генерирует договор подряда .docx из шаблона.
    """
    if not TEMPLATE_PATH.exists():
        raise FileNotFoundError(f"Шаблон не найден: {TEMPLATE_PATH}")

    doc = Document(str(TEMPLATE_PATH))

    # --- Данные ---
    area = parsed.get("area_m2", 0) or 0
    thickness = parsed.get("thickness_mm_avg", 0) or 0
    address = parsed.get("address") or "___"

    total = estimate["grand_total"]
    if include_sand_removal:
        total += 5000

    total_words = _num_to_words(total)
    total_formatted = f"{total:,}".replace(",", " ")

    full_name = client_data.get("full_name") or "___"
    passport_series = client_data.get("passport_series") or "____"
    passport_number = client_data.get("passport_number") or "______"
    passport = f"{passport_series} {passport_number}"
    issued_by = client_data.get("passport_issued_by") or "___"
    passport_date = client_data.get("passport_date") or "___"
    reg_address = client_data.get("registration_address") or "___"

    contract_num = client_data.get("contract_number") or "___"
    contract_date = client_data.get("contract_date") or ""
    if not contract_date:
        now = datetime.now(MSK)
        contract_date = f"{now.day:02d}.{now.month:02d}.{now.year}"

    # Дата для шапки
    try:
        cd_parts = contract_date.split(".")
        date_header = f"«{cd_parts[0]}» {MONTHS_RU[int(cd_parts[1])]} {cd_parts[2]} г."
    except Exception:
        date_header = contract_date

    work_start = client_data.get("work_start_date") or "___"
    work_end = client_data.get("work_end_date") or "___"
    payment_terms = client_data.get("payment_terms") or "___"

    # Фамилия И.О.
    name_parts = full_name.split()
    if len(name_parts) >= 3:
        short_name = f"{name_parts[0]} {name_parts[1][0]}.{name_parts[2][0]}."
    elif len(name_parts) == 2:
        short_name = f"{name_parts[0]} {name_parts[1][0]}."
    else:
        short_name = full_name

    # --- Замены в документе ---

    # Номер договора
    _replace_in_doc(doc, "№ 1/26ФЛ", f"№ {contract_num}/26ФЛ")
    _replace_in_doc(doc, "№1/26", f"№{contract_num}/26")

    # Дата в шапке
    _replace_in_doc(doc, "«  » января 2026 г.", date_header)

    # ФИО клиента (в шапке после "и")
    _replace_in_doc(doc, ", именуемая в дальнейшем", f" {full_name}, именуемый(ая) в дальнейшем")

    # Адрес объекта
    _replace_in_doc(doc, "г. Нижний Новгород, ул. Норильская, д. 16, кв. 5", address)

    # Площадь п.1.2.1 — ищем underlined "м2"
    _replace_in_doc(doc, "составляет м2", f"составляет {area} м2")

    # Толщина п.1.2.1
    ker_data = parsed.get("keramzit") or {}
    ker_thick = ker_data.get("thickness_mm", 0)

    if ker_thick:
        ker_area = ker_data.get("area_m2", 0)
        # С керамзитом: площадь и толщина керамзитного основания
        _replace_in_doc(doc, "составляет  мм",
                        f"составляет {thickness} мм.\n"
                        f"Площадь керамзитного основания составляет {ker_area} м2, "
                        f"средняя толщина керамзитного основания составляет {ker_thick}мм")
        _replace_in_doc(doc, "составляет мм",
                        f"составляет {thickness} мм.\n"
                        f"Площадь керамзитного основания составляет {ker_area} м2, "
                        f"средняя толщина керамзитного основания составляет {ker_thick}мм")
    else:
        _replace_in_doc(doc, "составляет  мм", f"составляет {thickness} мм")
        _replace_in_doc(doc, "составляет мм", f"составляет {thickness} мм")

    # Сумма п.2.1
    _replace_in_doc(doc, "(тысяч) рублей 00 копеек",
                    f"{total_formatted} ({total_words}) рублей 00 копеек")

    # --- Вставка таблицы сметы после п.2.1 ---
    _insert_estimate_table(doc, parsed, estimate, area, thickness, grade, include_sand_removal)

    # Толщина п.2.5
    if ker_thick:
        _replace_in_doc(doc, "средней толщине стяжки  мм",
                        f"средней толщине стяжки {thickness}мм и средней толщины керамзитного основания {ker_thick}мм")
        _replace_in_doc(doc, "средней толщине стяжки мм",
                        f"средней толщине стяжки {thickness}мм и средней толщины керамзитного основания {ker_thick}мм")
    else:
        _replace_in_doc(doc, "средней толщине стяжки  мм", f"средней толщине стяжки {thickness} мм")
        _replace_in_doc(doc, "средней толщине стяжки мм", f"средней толщине стяжки {thickness} мм")

    # Дата начала работ п.3.1 и завершения п.3.2
    # Текст разбит по runs — заменяем напрямую
    for para in doc.paragraphs:
        if "начать работы на объекте Заказчика" in para.text:
            # Собираем весь текст, заменяем, пишем в первый run
            new_text = f"3.1. Подрядчик обязуется начать работы на объекте Заказчика {work_start}"
            para.runs[0].text = new_text
            for r in para.runs[1:]:
                r.text = ""
            break

    for para in doc.paragraphs:
        if "готовности сдачи результата работ" in para.text and "3.2" in para.text:
            new_text = f"3.2. Подрядчик обязуется завершить работы и сообщить Заказчику о готовности сдачи результата работ {work_end}"
            para.runs[0].text = new_text
            for r in para.runs[1:]:
                r.text = ""
            break

    # Дата в приложении
    _replace_in_doc(doc, "от 13.01.2026 г.", f"от {contract_date} г.")

    # График финансирования — свободный текст условий оплаты
    _replace_in_doc(doc, "16.01.2025г Расчет              рублей.",
                    payment_terms)
    _replace_in_doc(doc, "16.01.2025г Расчет", payment_terms)

    # Реквизиты заказчика (в таблицах)
    _replace_in_doc(doc, "ФИО: ", f"ФИО: {full_name}")
    _replace_in_doc(doc, "Паспорт: ", f"Паспорт: {passport}")
    _replace_in_doc(doc, "Выдан: ", f"Выдан: {issued_by}")
    _replace_in_doc(doc, "Дата выдачи: ", f"Дата выдачи: {passport_date}\n\nЗарегистрирован по адресу:\n{reg_address}")

    # Подпись
    _replace_in_doc(doc, "( )", f"({short_name})")

    # --- Разрыв страницы перед Приложением 1 ---
    from docx.oxml.ns import qn as _qn
    for i, para in enumerate(doc.paragraphs):
        if "Приложение 1" in para.text and "ДОГОВОРУ" in para.text:
            # Добавляем разрыв страницы перед этим параграфом
            run = para.runs[0] if para.runs else para.add_run()
            br = run._r.makeelement(_qn("w:br"), {_qn("w:type"): "page"})
            run._r.insert(0, br)
            break

    # --- Сохранение ---
    if output_path is None:
        name = full_name.split()[0] if full_name != "___" else "клиент"
        output_path = f"/tmp/Договор_{contract_num}_{name}.docx"

    doc.save(output_path)
    logger.info("Договор сохранён: %s", output_path)
    return output_path


# ---------- Тест ----------

if __name__ == "__main__":
    from calculator import calculate_estimate

    parsed = {
        "area_m2": 36.9,
        "thickness_mm_avg": 86,
        "address": "г. Нижний Новгород, ул. Политбойцов, д.19, кв. 19",
    }

    estimate = calculate_estimate(
        area_m2=36.9, thickness_mm=86, is_city=True,
        grade="М150", floor=1,
    )

    client_data = {
        "full_name": "Котелков Виктор Михайлович",
        "passport_series": "2221",
        "passport_number": "295591",
        "passport_issued_by": "ГУ МВД России по Нижегородской области",
        "passport_date": "02.06.2021",
        "registration_address": "г. Нижний Новгород, ул. Политбойцов, д.19, кв. 19",
        "contract_number": "48",
        "contract_date": "04.03.2026",
        "work_start_date": "05.03.2026г.",
        "work_end_date": "05.03.2026г.",
        "payment_terms": "Аванс 26.03.2026 - 50000 руб. Окончательный расчет 28.03.2026 - 26517 руб.",
    }

    path = generate_contract(
        parsed=parsed,
        estimate=estimate,
        client_data=client_data,
        output_path="/home/claude/test_contract3.docx",
    )
    print(f"Договор: {path}")
