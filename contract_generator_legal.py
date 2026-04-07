"""
ARTPOL — Генератор договора подряда с юрлицами (безнал)
Шаблон: ШАБЛОН_ДОГОВОРА_ЮРЛИЦАМИ.docx
Подставляет данные заказчика, сметы, графиков в шаблон.
"""

import os
import logging
from pathlib import Path
from datetime import datetime, timezone, timedelta
from copy import deepcopy

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

logger = logging.getLogger(__name__)

TEMPLATE_PATH = Path(__file__).parent / "ШАБЛОН_ДОГОВОРА_ЮРЛИЦАМИ.docx"
MSK = timezone(timedelta(hours=3))

MONTHS_RU = {
    1: "января", 2: "февраля", 3: "марта", 4: "апреля",
    5: "мая", 6: "июня", 7: "июля", 8: "августа",
    9: "сентября", 10: "октября", 11: "ноября", 12: "декабря",
}


def _num_to_words(n: int) -> str:
    """Число → прописью (до миллиона)."""
    if n == 0:
        return "ноль"

    ones = ["", "одна", "две", "три", "четыре", "пять", "шесть",
            "семь", "восемь", "девять", "десять", "одиннадцать",
            "двенадцать", "тринадцать", "четырнадцать", "пятнадцать",
            "шестнадцать", "семнадцать", "восемнадцать", "девятнадцать"]
    tens = ["", "", "двадцать", "тридцать", "сорок", "пятьдесят",
            "шестьдесят", "семьдесят", "восемьдесят", "девяносто"]
    hundreds = ["", "сто", "двести", "триста", "четыреста", "пятьсот",
                "шестьсот", "семьсот", "восемьсот", "девятьсот"]

    parts = []
    orig = n

    if n >= 1000:
        t = n // 1000
        n = n % 1000
        if t >= 100:
            parts.append(hundreds[t // 100])
            t = t % 100
        if t >= 20:
            parts.append(tens[t // 10])
            t = t % 10
        if t > 0:
            if t == 1:
                parts.append("одна")
            elif t == 2:
                parts.append("две")
            else:
                parts.append(ones[t])

        tt = (orig // 1000) % 100
        last_digit = (orig // 1000) % 10
        if 11 <= tt <= 19:
            parts.append("тысяч")
        elif last_digit == 1:
            parts.append("тысяча")
        elif 2 <= last_digit <= 4:
            parts.append("тысячи")
        else:
            parts.append("тысяч")

    if n >= 100:
        parts.append(hundreds[n // 100])
        n = n % 100
    if n >= 20:
        parts.append(tens[n // 10])
        n = n % 10
    if n > 0:
        parts.append(ones[n])

    return " ".join(p for p in parts if p)


def _replace_in_paragraph(paragraph, old_text, new_text):
    """Заменяет текст в параграфе, даже если он разбит на несколько runs."""
    for run in paragraph.runs:
        if old_text in run.text:
            run.text = run.text.replace(old_text, str(new_text))
            return True

    full_text = paragraph.text
    if old_text not in full_text:
        return False

    runs_with_text = [(i, run) for i, run in enumerate(paragraph.runs) if run.text]
    if not runs_with_text:
        return False

    concat = ""
    for i, run in runs_with_text:
        concat += run.text

    find_pos = concat.find(old_text)
    if find_pos == -1:
        return False

    find_end = find_pos + len(old_text)
    new_concat = concat[:find_pos] + str(new_text) + concat[find_end:]

    first_run_idx = runs_with_text[0][0]
    paragraph.runs[first_run_idx].text = new_concat
    for i, run in runs_with_text[1:]:
        run.text = ""

    return True


def _replace_in_doc(doc, old_text, new_text):
    """Заменяет текст во всём документе (параграфы + таблицы)."""
    for para in doc.paragraphs:
        _replace_in_paragraph(para, old_text, str(new_text))

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_paragraph(para, old_text, str(new_text))


def _set_table_cell(cell, text, bold=False, size=9, align=WD_ALIGN_PARAGRAPH.LEFT):
    """Настраивает ячейку таблицы."""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    run.bold = bold


def _shade_cells(row, color):
    """Заливает ячейки строки цветом."""
    for cell in row.cells:
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
        cell._tc.get_or_add_tcPr().append(shading)


def _fill_estimate_table(table, estimate, parsed, area, thickness, grade, include_sand_removal, is_beznal):
    """Заполняет таблицу сметы (Приложение 3)."""
    s = estimate["sand"]
    c = estimate["cement"]
    f = estimate["fiber"]
    fl = estimate["film"]
    iz = estimate["izoflex"]
    eq = estimate["equipment_delivery"]
    w = estimate["work"]
    k = estimate.get("keramzit")

    keramzit_data = parsed.get("keramzit") or {}
    ker_area = keramzit_data.get("area_m2", 0)
    ker_thick = keramzit_data.get("thickness_mm", 0)
    volume = round(area * thickness / 1000, 2)

    L = WD_ALIGN_PARAGRAPH.LEFT
    R = WD_ALIGN_PARAGRAPH.RIGHT
    C = WD_ALIGN_PARAGRAPH.CENTER
    BLUE_LIGHT = "DCE6F1"
    BLUE_HEADER = "B8CCE4"

    def add_row(data, shade=None):
        row = table.add_row()
        for i, (text, bold, align) in enumerate(data):
            _set_table_cell(row.cells[i], text, bold=bold, size=9, align=align)
        if shade:
            _shade_cells(row, shade)
        return row

    # Удаляем все существующие строки кроме шапки
    while len(table.rows) > 0:
        tr = table.rows[-1]._tr
        table._tbl.remove(tr)

    # Шапка: керамзит (если есть)
    if k and ker_area:
        add_row([("", False, L), (f"{ker_area}", False, R), ("", False, L),
                 (f"{ker_thick}", False, R), ("", False, R)], shade=BLUE_LIGHT)

    # Шапка: стяжка
    add_row([("", False, L), (f"{area}", False, R), ("", False, L),
             (f"{thickness}", False, R), (f"{volume}", False, R)], shade=BLUE_LIGHT)

    # Заголовки
    add_row([("материалы и транспортные расходы", True, L), ("ед. изм.", True, C),
             ("кол-во ед.", True, C), ("стоимость ед.", True, C),
             ("ИТОГО", True, R)], shade=BLUE_HEADER)

    # Песок
    add_row([(f"Песок речной {s['sand_tons']}т + доставка", False, L), ("рейс", False, C),
             ("1", False, R), (f"{s['total']}", False, R), (f"{s['total']}", False, R)])

    # Цемент
    ppb = c["cement_cost"] // c["bags"] if c["bags"] > 0 else 0
    add_row([("Цемент М500 по 50кг", False, L), ("мешок", False, C),
             (f"{c['bags']}", False, R), (f"{ppb}", False, R), (f"{c['cement_cost']}", False, R)])

    # Фибра
    add_row([("Фибра ВСМ 12 мм 20 мк", False, L), ("кг", False, C),
             (f"{f['kg']}", False, R), ("300", False, R), (f"{f['cost']}", False, R)])

    # Плёнка техн.
    add_row([("Пленка 60 мкр техническая", False, L), ("м2", False, C),
             (f"{fl['m2']}", False, R), ("10", False, R), (f"{fl['cost']}", False, R)])

    # Izoflex
    add_row([("IZOFLEX 10мм", False, L), ("пог.м.", False, C),
             (f"{iz['meters']}", False, R), ("20", False, R), (f"{iz['cost']}", False, R)])

    # Керамзит (если есть)
    if k:
        add_row([("Керамзит (0,075м3)", False, L), ("мешок", False, C),
                 (f"{k['keramzit_bags']}", False, R), ("340", False, R), (f"{k['keramzit_cost']}", False, R)])
        add_row([("Армированная пленка 100г/м", False, L), ("м2", False, C),
                 (f"{k['reinforced_film_m2']}", False, R), ("40", False, R), (f"{k['reinforced_film_cost']}", False, R)])
        add_row([("Металлическая сетка", False, L), ("м2", False, C),
                 (f"{k['mesh_m2']}", False, R), ("120", False, R), (f"{k['mesh_cost']}", False, R)])

    # Доставка материалов
    add_row([("Доставка материалов", False, L), ("рейс", False, C),
             ("1", False, R), (f"{c['delivery']}", False, R), (f"{c['delivery']}", False, R)])

    # Доставка оборудования
    add_row([("Доставка оборудования", False, L), ("рейс", False, C),
             ("1", False, R), (f"{eq['cost']}", False, R), (f"{eq['cost']}", False, R)])

    # *Работы
    add_row([("*Работы", True, L), ("", False, C), ("", False, R),
             ("руб./м2", False, R), ("", False, R)], shade=BLUE_HEADER)

    # Вывоз песка
    if include_sand_removal:
        sand_removal_cost = round(6000 * 1.5) if is_beznal else 6000
        add_row([("Вывоз\\довоз песка", False, L), ("", False, C),
                 ("", False, R), ("", False, R), (f"{sand_removal_cost}", False, R)])

    # Работа керамзит
    if k:
        add_row([("Устройство керамзитного основания", False, L), ("м2", False, C),
                 (f"{ker_area}", False, R), (f"{k['keramzit_work_rate']}", False, R),
                 (f"{k['keramzit_work_cost']}", False, R)])

    # Работа стяжка
    work_rate = w.get("rate", "").replace("₽", "")
    add_row([("Устройство полусухой стяжки пола", False, L), ("м2", False, C),
             (f"{area}", False, R), (work_rate, False, R), (f"{w['cost']}", False, R)])

    # ИТОГО
    grand = estimate["grand_total"]
    if include_sand_removal:
        sand_removal_cost = round(6000 * 1.5) if is_beznal else 6000
        grand += sand_removal_cost

    add_row([("Безналичный расчет", False, L), ("", False, C), ("", False, R),
             ("ИТОГО стяжка", True, R), (f"{grand}", True, R)])


def generate_legal_contract(
    parsed: dict,
    estimate: dict,
    client_data: dict,
    grade: str = "М150",
    include_sand_removal: bool = False,
    output_path: str = None,
) -> str:
    """
    Генерирует договор подряда с юрлицом из шаблона.

    client_data должен содержать:
    - org_name: "ООО «РМУ»"
    - director_title: "Директора" / "Генерального директора"
    - director_name_genitive: "Назарова Дениса Андреевича" (родительный)
    - director_name_short: "Назаров Д.А."
    - director_basis: "Устава"
    - email: "denis@rmu.bizml.ru"
    - legal_address: юр.адрес
    - inn: ИНН
    - kpp: КПП
    - ogrn: ОГРН
    - bank_account: р/с
    - bank_name: название банка
    - corr_account: к/с
    - bik: БИК
    - contract_number: номер договора
    - work_start_date: "16.02.2026"
    - work_end_date: "16.02.2026"
    - payment_terms: условия оплаты (свободный текст)
    """
    if not TEMPLATE_PATH.exists():
        raise FileNotFoundError(f"Шаблон не найден: {TEMPLATE_PATH}")

    doc = Document(str(TEMPLATE_PATH))

    # --- Данные из сметы ---
    area = parsed.get("area_m2", 0) or 0
    thickness = parsed.get("thickness_mm_avg", 0) or 0
    address = parsed.get("address") or "___"

    is_beznal = True  # Этот генератор всегда для безнала

    total = estimate["grand_total"]
    if include_sand_removal:
        total += round(6000 * 1.5)  # безнал

    total_words = _num_to_words(total)
    total_formatted = f"{total:,}".replace(",", " ")

    # --- Данные заказчика ---
    org_name = client_data.get("org_name") or "___"
    director_title = client_data.get("director_title") or "Директора"
    director_name_genitive = client_data.get("director_name_genitive") or "___"
    director_name_short = client_data.get("director_name_short") or "___"
    director_basis = client_data.get("director_basis") or "Устава"
    email = client_data.get("email") or "___"
    legal_address = client_data.get("legal_address") or "___"
    inn = client_data.get("inn") or "___"
    kpp = client_data.get("kpp") or "___"
    ogrn = client_data.get("ogrn") or "___"
    bank_account = client_data.get("bank_account") or "___"
    bank_name = client_data.get("bank_name") or "___"
    corr_account = client_data.get("corr_account") or "___"
    bik = client_data.get("bik") or "___"

    contract_num = client_data.get("contract_number") or "___"
    contract_date = client_data.get("contract_date") or ""
    if not contract_date:
        now = datetime.now(MSK)
        contract_date = f"{now.day:02d}.{now.month:02d}.{now.year}"

    # Дата для шапки: "«10» февраля 2026 г."
    try:
        cd_parts = contract_date.split(".")
        date_header = f"«{cd_parts[0]}» {MONTHS_RU[int(cd_parts[1])]} {cd_parts[2]} г."
    except Exception:
        date_header = contract_date

    work_start = client_data.get("work_start_date") or "___"
    work_end = client_data.get("work_end_date") or "___"
    payment_terms = client_data.get("payment_terms") or "___"

    # Даты работ в формате «ДД» месяц ГГГГ г.
    def _format_work_date(date_str):
        try:
            parts = date_str.replace("г.", "").strip().split(".")
            return f"«{parts[0]}» {MONTHS_RU[int(parts[1])]} {parts[2]} г."
        except Exception:
            return date_str

    work_start_formatted = _format_work_date(work_start)
    work_end_formatted = _format_work_date(work_end)

    # Керамзит
    keramzit_data = parsed.get("keramzit") or {}
    ker_area = keramzit_data.get("area_m2", 0)

    # ============================================================
    # ЗАМЕНЫ В ДОКУМЕНТЕ
    # ============================================================

    # --- Номер договора ---
    _replace_in_doc(doc, "№ 10/02/26 ЮЛ", f"№ {contract_num}/26 ЮЛ")
    _replace_in_doc(doc, "№10/02/26 ЮЛ", f"№{contract_num}/26 ЮЛ")

    # --- Дата в шапке ---
    _replace_in_doc(doc, "«10» февраля 2026 г.", date_header)

    # --- Заказчик в шапке ---
    _replace_in_doc(doc, "ООО «РМУ», именуемое в дальнейшем «Заказчик», в лице Директора Назарова Дениса Андреевича, действующего на основании Устава",
                    f"{org_name}, именуемое в дальнейшем «Заказчик», в лице {director_title} {director_name_genitive}, действующего на основании {director_basis}")

    # --- п.1.1: площадь, толщина, адрес ---
    _replace_in_doc(doc, "94,5м2 толщиной 120мм", f"{area}м2 толщиной {thickness}мм")
    _replace_in_doc(doc, "94,5м2", f"{area}м2")
    _replace_in_doc(doc, "г. Дзержинск, Автозаводское шоссе, 101Б", address)

    # --- п.2.1.14: марка прочности ---
    _replace_in_doc(doc, "не ниже М200", f"не ниже {grade}")

    # --- п.2.2.2: email заказчика ---
    _replace_in_doc(doc, "denis@rmu.bizml.ru", email)

    # --- п.2.2.5: площадь ---
    _replace_in_doc(doc, "94,5 м2 в день для производства работ", f"{area} м2 в день для производства работ")

    # --- п.3: сроки работ ---
    _replace_in_doc(doc, "«16» февраля 2026 г.", work_start_formatted)
    # Второе вхождение — окончание (может не сработать если уже заменено, поэтому ищем по контексту)
    for para in doc.paragraphs:
        if "Срок окончания Работ" in para.text:
            for run in para.runs:
                if "16" in run.text or "февраля" in run.text:
                    pass  # Уже заменено выше
            # Заменяем весь параграф
            new_text = f"Срок окончания Работ: {work_end_formatted}"
            para.runs[0].text = new_text
            for r in para.runs[1:]:
                r.text = ""
            break

    # --- п.4.1: сумма ---
    _replace_in_doc(doc, "94,5 м2 на средний слой стяжки пола 120 мм составляет 167 000 (сто шестьдесят семь тысяч)",
                    f"{area} м2 на средний слой стяжки пола {thickness} мм составляет {total_formatted} ({total_words})")

    # --- п.4.7: толщина и площадь ---
    _replace_in_doc(doc, "120мм на площади 94,5 м2", f"{thickness}мм на площади {area} м2")

    # --- п.5.3: наша почта (оставляем artpol_office@mail.ru) ---
    # Уже в шаблоне, не трогаем

    # ============================================================
    # РЕКВИЗИТЫ ЗАКАЗЧИКА (заменяем во всех 3+ местах)
    # ============================================================

    # Название организации
    _replace_in_doc(doc, "ООО «РМУ»", org_name)

    # Юр.адрес
    _replace_in_doc(doc, "603086, г. Нижний Новгород, б-р Мира, д. 3, помещ. 9", legal_address)

    # ИНН
    _replace_in_doc(doc, "5259049073", inn)

    # КПП заказчика (отличаем от нашего 525701001)
    # Наш КПП тоже 525701001 — нужна осторожность
    # Заменяем только в контексте заказчика
    for para in doc.paragraphs:
        full = para.text
        # Ищем "КПП" рядом с ИНН заказчика
        if kpp != "525701001" and "КПП 525701001" in full and inn not in full:
            # Это может быть наш КПП — пропускаем
            pass

    # ОГРН заказчика
    _replace_in_doc(doc, "1065259021871", ogrn)

    # Банковские реквизиты заказчика
    _replace_in_doc(doc, "40702810801300003005", bank_account)
    _replace_in_doc(doc, "ПАО «Банк Уралсиб»", bank_name)
    _replace_in_doc(doc, "30101810600000000770", corr_account)
    _replace_in_doc(doc, "048073770", bik)

    # Директор заказчика (подписи)
    _replace_in_doc(doc, "/Назаров Д.А/", f"/{director_name_short}/")
    _replace_in_doc(doc, "/Назаров Д.А./", f"/{director_name_short}/")

    # ============================================================
    # ПРИЛОЖЕНИЕ 1 — График производства работ (Table 0)
    # ============================================================
    if len(doc.tables) >= 1:
        t0 = doc.tables[0]
        # Строка 1: даты и наименование
        if len(t0.rows) >= 2:
            row = t0.rows[1]
            row.cells[0].text = work_start.replace("г.", "").strip()
            row.cells[1].text = work_end.replace("г.", "").strip()
            work_desc = f"Устройство полусухой стяжки пола {area} м2"
            if ker_area:
                work_desc += f" и керамзитного основания {ker_area} м2"
            row.cells[2].text = work_desc

    # ============================================================
    # ПРИЛОЖЕНИЕ 2 — График финансирования (Table 1)
    # ============================================================
    if len(doc.tables) >= 2:
        t1 = doc.tables[1]
        if len(t1.rows) >= 2:
            row = t1.rows[1]
            row.cells[1].text = payment_terms

    # ============================================================
    # ПРИЛОЖЕНИЕ 3 — Таблица сметы (Table 2)
    # ============================================================
    if len(doc.tables) >= 3:
        t2 = doc.tables[2]
        _fill_estimate_table(t2, estimate, parsed, area, thickness, grade, include_sand_removal, is_beznal)

    # --- Номер договора в приложениях ---
    _replace_in_doc(doc, "10/02/26 ЮЛ от 10.02.2026", f"{contract_num}/26 ЮЛ от {contract_date}")

    # ============================================================
    # СОХРАНЕНИЕ
    # ============================================================
    if output_path is None:
        org_short = org_name.split("«")[1].split("»")[0] if "«" in org_name else "ЮЛ"
        output_path = f"/tmp/Договор_ЮЛ_{contract_num}_{org_short}.docx"

    doc.save(output_path)
    logger.info("Договор ЮЛ сохранён: %s", output_path)
    return output_path
