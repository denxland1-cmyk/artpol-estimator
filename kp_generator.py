"""
ARTPOL — Генератор КП (коммерческое предложение)
Создаёт .docx файл из данных парсера + калькулятора.
Дизайн максимально приближен к оригинальному шаблону.
"""

import logging
from pathlib import Path
from datetime import datetime, timezone, timedelta

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

logger = logging.getLogger(__name__)

LOGO_PATH = Path(__file__).parent / "logo.png"
MSK = timezone(timedelta(hours=3))

# Ширины столбцов — сумма = 17.5 см (A4 минус поля 2+1.5)
COL_W = [Cm(6.5), Cm(1.8), Cm(2.0), Cm(3.2), Cm(4.0)]
TABLE_W = sum(c for c in COL_W)

GRAY_LIGHT = "DCE6F1"   # шапка таблицы (площадь/толщина/объём)
GRAY_HEADER = "B8CCE4"  # заголовки (материалы, *Работы)
WHITE = "FFFFFF"


def _shade_row(row, color):
    """Заливает все ячейки строки цветом."""
    for cell in row.cells:
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
        cell._tc.get_or_add_tcPr().append(shading)


def _set_cell(cell, text, bold=False, size=10, align=WD_ALIGN_PARAGRAPH.LEFT):
    """Настраивает текст ячейки с форматированием."""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = Pt(13)
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    run.bold = bold
    # Устанавливаем шрифт для кириллицы
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.makeelement(qn("w:rFonts"), {})
    rFonts.set(qn("w:ascii"), "Calibri")
    rFonts.set(qn("w:hAnsi"), "Calibri")
    rFonts.set(qn("w:cs"), "Calibri")
    rPr.insert(0, rFonts)


def _add_row(table, col_data, shade=None):
    """
    Добавляет строку. col_data: [(text, bold, align), ...]
    shade: цвет заливки или None
    """
    row = table.add_row()
    for i, (text, bold, align) in enumerate(col_data):
        _set_cell(row.cells[i], text, bold=bold, size=10, align=align)
    if shade:
        _shade_row(row, shade)
    return row


def _add_para(doc, text, bold=False, size=11, align=WD_ALIGN_PARAGRAPH.LEFT,
              space_before=0, space_after=2, color=None):
    """Добавляет параграф с настройками."""
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = Pt(14)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    run.bold = bold
    if color:
        run.font.color.rgb = color
    return p


def generate_kp(
    parsed: dict,
    estimate: dict,
    grade: str = "М150",
    payment_type: str = "наличными",
    include_sand_removal: bool = False,
    output_path: str = None,
) -> str:
    """
    Генерирует .docx файл КП.

    Возвращает путь к созданному файлу.
    """
    doc = Document()

    # --- Страница ---
    section = doc.sections[0]
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(1.5)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(0)

    L = WD_ALIGN_PARAGRAPH.LEFT
    R = WD_ALIGN_PARAGRAPH.RIGHT
    C = WD_ALIGN_PARAGRAPH.CENTER

    # --- Логотип ---
    if LOGO_PATH.exists():
        p = doc.add_paragraph()
        p.alignment = C
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run()
        run.add_picture(str(LOGO_PATH), width=Inches(2.8))

    # Подзаголовок
    _add_para(doc, "Полусухая стяжка", bold=True, size=14, align=C,
              space_before=2, space_after=0)
    _add_para(doc, "Механизированная штукатурка", bold=True, size=14, align=C,
              space_before=0, space_after=4)

    # --- Разделительная линия ---
    p_line = doc.add_paragraph()
    p_line.paragraph_format.space_before = Pt(2)
    p_line.paragraph_format.space_after = Pt(10)
    pPr = p_line._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="6" w:space="1" w:color="000000"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)

    # --- Данные объекта ---
    area = parsed.get("area_m2", 0)
    thickness = parsed.get("thickness_mm_avg", 0)
    address = parsed.get("address", "—")
    client_name = parsed.get("client_name", "")
    client_phone = parsed.get("client_phone", "")
    contact = f"{client_phone} {client_name}".strip() or "—"

    _add_para(doc, "Коммерческое предложение на устройство полусухой стяжки пола",
              size=11, space_before=0, space_after=1)

    # Дата
    date_str = datetime.now(MSK).strftime("%d.%m.%Y")
    _add_para(doc, f"от {date_str}", size=10, align=R,
              space_before=0, space_after=1, color=RGBColor(0x66, 0x66, 0x66))

    _add_para(doc, f"на площади {area} м2 средней толщиной {thickness} мм",
              bold=True, size=11, space_after=1)

    # Керамзит в шапке
    keramzit = estimate.get("keramzit")
    keramzit_data = parsed.get("keramzit") or {}
    ker_area = keramzit_data.get("area_m2", 0)
    ker_thick = keramzit_data.get("thickness_mm", 0)
    if keramzit and ker_area:
        _add_para(doc, "на устройство керамзитного основания", size=11, space_after=1)
        _add_para(doc, f"на площади {ker_area} м2 средней толщиной {ker_thick} мм",
                  bold=True, size=11, space_after=1)

    _add_para(doc, f"по адресу: {address}", size=11, space_after=1)
    _add_para(doc, f"Контактное лицо: {contact}", size=11, space_after=8)

    # --- Марочная прочность ---
    p_grade = doc.add_paragraph()
    p_grade.alignment = R
    p_grade.paragraph_format.space_after = Pt(4)
    r1 = p_grade.add_run("Марочная прочность ")
    r1.font.size = Pt(11)
    r1.font.name = "Calibri"
    r2 = p_grade.add_run(grade)
    r2.font.size = Pt(11)
    r2.font.name = "Calibri"
    r2.bold = True

    # --- ТАБЛИЦА СМЕТЫ ---
    s = estimate["sand"]
    c = estimate["cement"]
    f = estimate["fiber"]
    fl = estimate["film"]
    iz = estimate["izoflex"]
    eq = estimate["equipment_delivery"]
    w = estimate["work"]
    k = estimate.get("keramzit")
    volume = round(area * thickness / 1000, 3)

    table = doc.add_table(rows=0, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False

    # Устанавливаем ширину столбцов
    for i, width in enumerate(COL_W):
        table.columns[i].width = width

    # Шапка: керамзит (если есть) + стяжка
    if k and ker_area:
        _add_row(table, [
            ("", False, L),
            (f"{ker_area}", False, R),
            ("", False, L),
            (f"{ker_thick}", False, R),
            ("", False, R),
        ], shade=GRAY_LIGHT)

    _add_row(table, [
        ("", False, L),
        (f"{area}", False, R),
        ("", False, L),
        (f"{thickness}", False, R),
        (f"{volume}", False, R),
    ], shade=GRAY_LIGHT)

    # Заголовки — голубой фон
    _add_row(table, [
        ("материалы и транспортные расходы", True, L),
        ("ед. изм.", True, C),
        ("кол-во ед.", True, C),
        ("Стоимость ед.", True, C),
        ("ИТОГО", True, R),
    ], shade=GRAY_HEADER)

    # Песок
    _add_row(table, [
        (f"Песок {s['sand_tons']}т+доставка", False, L),
        ("рейс", False, C),
        ("1", False, R),
        (f"{s['total']}", False, R),
        (f"{s['total']}", False, R),
    ])

    # Цемент
    price_per_bag = c["cement_cost"] // c["bags"] if c["bags"] > 0 else 0
    _add_row(table, [
        ("Цемент М500 по 50 кг", False, L),
        ("мешок", False, C),
        (f"{c['bags']}", False, R),
        (f"{price_per_bag}", False, R),
        (f"{c['cement_cost']}", False, R),
    ])

    # Фибра
    _add_row(table, [
        ("Фибра ВСМ 12 мм 20 мк", False, L),
        ("кг", False, C),
        (f"{f['kg']}", False, R),
        ("300", False, R),
        (f"{f['cost']}", False, R),
    ])

    # Плёнка
    _add_row(table, [
        ("Пленка 60 мкр техническая", False, L),
        ("м2", False, C),
        (f"{fl['m2']}", False, R),
        ("10", False, R),
        (f"{fl['cost']}", False, R),
    ])

    # Izoflex
    _add_row(table, [
        ("IZOFLEX 10 мм", False, L),
        ("пог.м.", False, C),
        (f"{iz['meters']}", False, R),
        ("20", False, R),
        (f"{iz['cost']}", False, R),
    ])

    # Керамзит материалы (если есть)
    if k:
        _add_row(table, [
            ("Керамзит (0,075м3)", False, L),
            ("мешок", False, C),
            (f"{k['keramzit_bags']}", False, R),
            ("340", False, R),
            (f"{k['keramzit_cost']}", False, R),
        ])

        _add_row(table, [
            ("Армированная пленка 100г/м", False, L),
            ("м2", False, C),
            (f"{k['reinforced_film_m2']}", False, R),
            ("40", False, R),
            (f"{k['reinforced_film_cost']}", False, R),
        ])

        _add_row(table, [
            ("Металлическая сетка", False, L),
            ("м2", False, C),
            (f"{k['mesh_m2']}", False, R),
            ("120", False, R),
            (f"{k['mesh_cost']}", False, R),
        ])

    # Доставка материалов
    _add_row(table, [
        ("Доставка материалов", False, L),
        ("рейс", False, C),
        ("1", False, R),
        (f"{c['delivery']}", False, R),
        (f"{c['delivery']}", False, R),
    ])

    # Доставка оборудования
    _add_row(table, [
        ("Доставка/вывоз оборудования", False, L),
        ("рейс", False, C),
        ("1", False, R),
        (f"{eq['cost']}", False, R),
        (f"{eq['cost']}", False, R),
    ])

    # *Работы — голубой фон
    _add_row(table, [
        ("*Работы", True, L),
        ("", False, C),
        ("", False, R),
        ("руб./м2", False, R),
        ("", False, R),
    ], shade=GRAY_HEADER)

    # Вывоз песка (опционально)
    is_beznal = estimate.get("payment_type") == "безналичный расчет"
    sand_removal_cost = round(5000 * 1.5) if is_beznal else 5000
    if include_sand_removal:
        _add_row(table, [
            ("Вывоз\\довоз песка", False, L),
            ("", False, C),
            ("", False, R),
            ("", False, R),
            (f"{sand_removal_cost}", False, R),
        ])

    # Работа: керамзит (если есть)
    if k:
        _add_row(table, [
            ("Устройство керамзитного основания", False, L),
            ("м2", False, C),
            (f"{ker_area}", False, R),
            (f"{k['keramzit_work_rate']}", False, R),
            (f"{k['keramzit_work_cost']}", False, R),
        ])

    # Устройство стяжки
    work_rate = w.get("rate", "")
    # Убираем ₽ из rate для таблицы
    work_rate_clean = work_rate.replace("₽", "").strip()
    _add_row(table, [
        ("Устройство полусухой стяжки пола", False, L),
        ("м2", False, C),
        (f"{area}", False, R),
        (work_rate_clean, False, R),
        (f"{w['cost']}", False, R),
    ])

    # ИТОГО
    grand = estimate["grand_total"]
    if include_sand_removal:
        grand += sand_removal_cost

    _add_row(table, [
        ("", False, L),
        ("", False, C),
        ("", False, R),
        ("ИТОГО стяжка", True, R),
        (f"{grand}", True, R),
    ])

    # Стиль таблицы — рамки
    table.style = "Table Grid"

    # Принудительная ширина таблицы на всю страницу
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>')
    # Удаляем старый tblW если есть
    for old in tblPr.findall(qn("w:tblW")):
        tblPr.remove(old)
    tblPr.append(tblW)

    # Устанавливаем ширину ячеек в каждой строке
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = COL_W[i]
            # Вертикальное выравнивание по центру
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            vAlign = parse_xml(f'<w:vAlign {nsdecls("w")} w:val="center"/>')
            tcPr.append(vAlign)

    # --- Текст под таблицей ---

    # Особые условия
    special = parsed.get("special_conditions", [])
    if special:
        _add_para(doc, "Особые условия: " + ", ".join(special),
                  bold=False, size=10, space_before=6, space_after=4)

    if include_sand_removal:
        _add_para(doc, "! Вывоз строительного мусора не входит в стоимость",
                  bold=True, size=10, space_before=4, space_after=2)
    else:
        _add_para(doc, "! Вывоз остатков песка и строительного мусора не входит в стоимость",
                  bold=True, size=10, space_before=4, space_after=2)

    _add_para(doc, f"! Оплата производится {payment_type}",
              bold=True, size=10, space_after=2)

    if is_beznal:
        _add_para(doc, "! В стоимость включен НДС 22%",
                  bold=True, size=10, space_after=2)

    _add_para(doc,
        "! Мы дарим Вам возможность приобретать любые товары в любых "
        "строительных магазинах с Выгодой для Вас до 10% от общей суммы заказа!",
        bold=True, size=10, space_after=2)

    _add_para(doc, "Акция не имеет срока давности!", bold=True, size=10, space_after=2)

    _add_para(doc, "За подробностями обращайтесь к Вашему персональному менеджеру!",
              bold=True, size=10, space_after=4)

    # --- Партнёры ---
    _add_para(doc, "Наши партнеры:", bold=True, size=10, align=R, space_before=4, space_after=0)

    partners = [
        "ЛЕМАНА ПРО", "МАКСИДОМ", "ОРДЕР", "ОБИ", "САКСЭС",
        "СТРОЙРЕМ",
    ]
    for partner in partners:
        p = doc.add_paragraph()
        p.alignment = R
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = Pt(12)
        run = p.add_run(f"➤  {partner}")
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = "Calibri"

    # --- Сохранение ---
    if output_path is None:
        client = (client_name or "клиент").replace(" ", "_")
        output_path = f"/tmp/KP_{client}_{area}m2.docx"

    doc.save(output_path)
    logger.info("КП сохранено: %s", output_path)
    return output_path


# ---------- Быстрый тест ----------

if __name__ == "__main__":
    from calculator import calculate_estimate

    parsed = {
        "area_m2": 70.7,
        "thickness_mm_avg": 70,
        "address": "г. Нижний Новгород, ул. Белозерская, д. 3, кв. 86",
        "client_name": "Александр",
        "client_phone": "+79290448448",
    }

    estimate = calculate_estimate(
        area_m2=70.7, thickness_mm=70, is_city=True,
        grade="М150", floor=1,
    )

    path = generate_kp(
        parsed=parsed,
        estimate=estimate,
        grade="М150",
        payment_type="наличными",
        include_sand_removal=True,
        output_path="/home/claude/test_kp_v2.docx",
    )
    print(f"КП сохранено: {path}")
